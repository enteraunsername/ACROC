from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import re

BASE = Path(r'D:\pythonProject\ACROC')
md_path = BASE / '第四章工作报告与未来工作设想.md'
out_path = BASE / '第四章工作报告与未来工作设想.docx'
text = md_path.read_text(encoding='utf-8')

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x1F, 0x1F)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = 'F2F4F7'


def set_run_font(run, latin='Calibri', east_asia='宋体', size=None, color=None, bold=None, italic=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:eastAsia'), east_asia)
    rfonts.set(qn('w:cs'), latin)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, latin='Calibri', east_asia='宋体', size=None, color=None, bold=None):
    font = style.font
    font.name = latin
    if size is not None:
        font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:eastAsia'), east_asia)
    rfonts.set(qn('w:cs'), latin)


def set_paragraph_border_bottom(paragraph, color='2E74B5', sz='8', space='8'):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement('w:tblGrid')
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[idx]))
            tcW.set(qn('w:type'), 'dxa')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    return run


def add_rich_text(paragraph, content, default_size=11, default_color=INK):
    parts = re.split(r'(\*\*[^*]+\*\*)', content)
    for part in parts:
        if not part:
            continue
        bold = part.startswith('**') and part.endswith('**')
        txt = part[2:-2] if bold else part
        run = paragraph.add_run(txt)
        set_run_font(run, size=default_size, color=default_color, bold=bold)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
set_style_font(styles['Normal'], size=11, color=INK)
styles['Normal'].paragraph_format.space_after = Pt(6)
styles['Normal'].paragraph_format.line_spacing = 1.10

set_style_font(styles['Title'], east_asia='微软雅黑', size=22, color=RGBColor(0,0,0), bold=True)
styles['Title'].paragraph_format.space_after = Pt(8)

for name, size, color, before, after, east in [
    ('Heading 1', 16, BLUE, 16, 8, '微软雅黑'),
    ('Heading 2', 13, BLUE, 12, 6, '微软雅黑'),
    ('Heading 3', 12, DARK_BLUE, 8, 4, '微软雅黑'),
]:
    style = styles[name]
    set_style_font(style, east_asia=east, size=size, color=color, bold=True)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.10

# Custom callout style.
if 'Report Callout' not in styles:
    callout = styles.add_style('Report Callout', WD_STYLE_TYPE.PARAGRAPH)
else:
    callout = styles['Report Callout']
set_style_font(callout, size=10.5, color=RGBColor(0x33,0x33,0x33))
callout.paragraph_format.left_indent = Inches(0.15)
callout.paragraph_format.right_indent = Inches(0.1)
callout.paragraph_format.space_before = Pt(4)
callout.paragraph_format.space_after = Pt(10)
callout.paragraph_format.line_spacing = 1.10

# Header/footer.
header_p = section.header.paragraphs[0]
header_p.text = ''
r = header_p.add_run('第四章工作报告与未来工作设想')
set_run_font(r, size=9, color=MUTED)
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

footer_p = section.footer.paragraphs[0]
footer_p.text = ''
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer_p.add_run('第 ')
set_run_font(r, size=9, color=MUTED)
field_run = add_field(footer_p, 'PAGE')
set_run_font(field_run, size=9, color=MUTED)
r = footer_p.add_run(' 页')
set_run_font(r, size=9, color=MUTED)

# Cover/title block.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(2)
r = p.add_run('工作报告')
set_run_font(r, east_asia='微软雅黑', size=12, color=MUTED, bold=True)

title = doc.add_paragraph(style='Title')
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_rich_text(title, '第四章工作报告与未来工作设想', default_size=22, default_color=RGBColor(0,0,0))

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
r = subtitle.add_run('基于《基于化学反应优化的聚类算法研究》第四章“自适应化学反应优化聚类算法”整理')
set_run_font(r, east_asia='微软雅黑', size=11, color=MUTED)

meta = doc.add_table(rows=3, cols=2)
set_table_geometry(meta, [1440, 7920])
meta.style = 'Table Grid'
rows = [('文档类型', '正式工作报告'), ('内容范围', '第四章：自适应化学反应优化聚类算法'), ('整理日期', '2026年5月18日')]
for row, (label, value) in zip(meta.rows, rows):
    for cell in row.cells:
        cell.text = ''
    shade_cell(row.cells[0], LIGHT_FILL)
    pr = row.cells[0].paragraphs[0]
    pr.paragraph_format.space_after = Pt(0)
    rr = pr.add_run(label)
    set_run_font(rr, east_asia='微软雅黑', size=10.5, color=DARK_BLUE, bold=True)
    pv = row.cells[1].paragraphs[0]
    pv.paragraph_format.space_after = Pt(0)
    rv = pv.add_run(value)
    set_run_font(rv, size=10.5, color=INK)

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(14)
rule.paragraph_format.space_after = Pt(6)
set_paragraph_border_bottom(rule)

lead = doc.add_paragraph(style='Report Callout')
add_rich_text(lead, '摘要：本报告围绕第四章提出的 ACROC 算法展开，梳理其研究背景、目标、方法设计、实验验证与阶段性结论，并提出后续参数自适应、鲁棒性提升、工程效率优化和真实场景验证等未来工作方向。', default_size=10.5, default_color=RGBColor(0x33,0x33,0x33))

# Parse body, skipping top title duplicated on cover.
lines = text.splitlines()
current_para = []
first_title_skipped = False

def flush_para():
    global current_para
    if current_para:
        content = ''.join(current_para).strip()
        if content:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.10
            add_rich_text(p, content)
        current_para = []

for raw in lines:
    line = raw.rstrip()
    stripped = line.strip()
    if not stripped:
        flush_para()
        continue
    if stripped.startswith('# '):
        if not first_title_skipped:
            first_title_skipped = True
            continue
        flush_para()
        p = doc.add_paragraph(stripped[2:].strip(), style='Title')
        continue
    if stripped.startswith('## '):
        flush_para()
        p = doc.add_paragraph(stripped[3:].strip(), style='Heading 1')
        continue
    if stripped.startswith('### '):
        flush_para()
        p = doc.add_paragraph(stripped[4:].strip(), style='Heading 2')
        continue
    m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
    if m:
        flush_para()
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        add_rich_text(p, m.group(2), default_size=11, default_color=INK)
        continue
    # Treat indented markdown paragraphs as ordinary paragraphs.
    current_para.append(stripped)
flush_para()

# Final quality pass: avoid widows for headings and set body Chinese fonts explicitly.
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        p.paragraph_format.keep_with_next = True
    for run in p.runs:
        if not run.text:
            continue
        # Preserve title/header choices; make sure East Asian glyphs are declared.
        if p.style.name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
            east = '微软雅黑'
        else:
            east = '宋体'
        set_run_font(run, east_asia=east)

doc.core_properties.title = '第四章工作报告与未来工作设想'
doc.core_properties.subject = '基于化学反应优化的聚类算法研究第四章工作报告'
doc.core_properties.author = 'Codex'
doc.core_properties.keywords = 'ACROC, 聚类分析, 化学反应优化, 工作报告, 未来工作设想'

doc.save(out_path)
print(out_path)
